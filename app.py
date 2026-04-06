import streamlit as st
import json
import tempfile
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import fitz  # PyMuPDF
from groq import Groq

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Poetry OCR Extractor",
    page_icon="📖",
    layout="wide",
)

st.title("📖 Poetry OCR Extractor")
st.markdown(
    "Sube un PDF con poemas. Extrae el texto y usa **Groq (gratis)** para "
    "limpiar y estructurar el contenido, luego descarga un Word limpio."
)

# ── PDF helpers ───────────────────────────────────────────────────────────────

def extract_text_half(pdf_path: str, page_index: int, half: str) -> str:
    """
    Extract text from only the left or right half of a page.
    half: "left" | "right" | "full"
    """
    doc = fitz.open(pdf_path)
    page = doc[page_index]
    rect = page.rect  # full page bounding box

    if half == "left":
        clip = fitz.Rect(rect.x0, rect.y0, rect.x1 / 2, rect.y1)
    elif half == "right":
        clip = fitz.Rect(rect.x1 / 2, rect.y0, rect.x1, rect.y1)
    else:
        clip = rect

    text = page.get_text("text", clip=clip, sort=True)
    doc.close()
    return text.strip()


def is_text_empty(text: str) -> bool:
    """Return True if the extracted text has no meaningful content."""
    return len(text.replace("\n", "").replace(" ", "")) < 20


def detect_page_mode(pdf_path: str, page_index: int) -> str:
    """
    Detect if a page is a double spread or single.
    Compares the aspect ratio: landscape → likely double spread.
    Returns "double" or "single".
    """
    doc = fitz.open(pdf_path)
    page = doc[page_index]
    rect = page.rect
    doc.close()
    ratio = rect.width / rect.height
    # Landscape pages are likely double spreads
    return "double" if ratio > 1.2 else "single"


# ── Groq prompt ───────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """Eres un experto en edición y transcripción de poesía hispanohablante.
Recibirás el texto CRUDO extraído automáticamente de una página de un libro de poemas.
El texto está sucio: tiene números de línea a la izquierda, anotaciones editoriales
a la derecha como (GB¹, OC¹, LV², P³...), encabezados de página repetidos, y números de página.

Tu tarea es limpiar y estructurar ese texto.

Devuelve ÚNICAMENTE un objeto JSON válido con esta forma exacta
(sin bloques de código, sin texto antes o después):

{
  "blank": false,
  "page_header": "texto del encabezado de página si existe, o null",
  "poems": [
    {
      "title": "TÍTULO DEL POEMA en mayúsculas tal como aparece, o null",
      "sections": [
        {
          "speaker": "Nombre del hablante si aparece (ej: 'El paciente:', 'El médico:'), o null",
          "lines": ["verso 1", "verso 2", "..."]
        }
      ]
    }
  ],
  "footnotes": ["nota al pie 1", "nota al pie 2"]
}

Reglas CRÍTICAS:
1. Si la página no tiene poemas ni contenido relevante, devuelve {"blank": true, "page_header": null, "poems": [], "footnotes": []}.
2. Una sola página puede contener DOS o más poemas completos — inclúyelos TODOS en el array "poems".
3. Un poema puede tener múltiples secciones con distintos hablantes.
   Ejemplo: "El mal del siglo" tiene sección "El paciente:" y sección "El médico:".
4. Si no hay hablante en una sección, pon null en "speaker".
5. ELIMINA completamente:
   - Números de línea (ej: "5", "10", "15" solos al inicio de línea)
   - Anotaciones editoriales entre paréntesis: (GB¹), (OC¹), (LV²), (P³), [sin comas en OC¹], etc.
   - Números de página solos
   - Encabezados repetidos de página (ej: "GOTAS AMARGAS", "JOSÉ ASUNCIÓN SILVA")
6. Conserva tildes, puntos suspensivos, signos de exclamación/interrogación, mayúsculas originales.
7. Las notas al pie (letra pequeña al fondo, generalmente empiezan con superíndice) van en "footnotes".
8. Los títulos de poema suelen estar centrados y en mayúsculas.
9. Los nombres de hablantes suelen terminar en dos puntos (ej: "El paciente:").
"""


def structure_page(client: Groq, raw_text: str, label: str) -> dict:
    """Send raw text to Groq and get back structured JSON."""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    f"Texto crudo de {label}:\n\n"
                    f"```\n{raw_text}\n```\n\n"
                    "Limpia y estructura. Devuelve solo el JSON."
                ),
            },
        ],
        temperature=0.1,
        max_tokens=2048,
    )
    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


# ── Word document builder ─────────────────────────────────────────────────────

def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def build_docx(all_pages: list) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(11)

    seen_titles: set = set()
    first_content = True

    for page_data in all_pages:
        if page_data.get("blank"):
            if not first_content:
                add_page_break(doc)
            else:
                doc.add_paragraph()
            continue

        poems = page_data.get("poems", [])
        footnotes = page_data.get("footnotes", [])

        if not poems and not footnotes:
            continue

        first_content = False

        for poem in poems:
            title = poem.get("title") or ""
            title_key = title.strip().upper()

            if title_key and title_key not in seen_titles:
                seen_titles.add(title_key)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(title.strip())
                run.bold = True
                run.font.size = Pt(13)
                p.paragraph_format.space_before = Pt(20)
                p.paragraph_format.space_after = Pt(8)

            for section in poem.get("sections", []):
                speaker = section.get("speaker")
                if speaker:
                    sp = doc.add_paragraph()
                    sr = sp.add_run(speaker.strip())
                    sr.italic = True
                    sr.font.size = Pt(11)
                    sp.paragraph_format.space_before = Pt(8)
                    sp.paragraph_format.space_after = Pt(2)

                for line in section.get("lines", []):
                    lp = doc.add_paragraph()
                    lp.add_run(line)
                    lp.paragraph_format.space_before = Pt(0)
                    lp.paragraph_format.space_after = Pt(1)
                    lp.paragraph_format.left_indent = Inches(0.5)

            gap = doc.add_paragraph()
            gap.paragraph_format.space_after = Pt(10)

        for fn in footnotes:
            fp = doc.add_paragraph()
            fr = fp.add_run(fn.strip())
            fr.font.size = Pt(8.5)
            fr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            fp.paragraph_format.space_before = Pt(0)
            fp.paragraph_format.space_after = Pt(1)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    with open(tmp_path, "rb") as f:
        data = f.read()
    os.unlink(tmp_path)
    return data


# ── Page range parser ─────────────────────────────────────────────────────────

def parse_page_range(spec: str, total: int) -> list:
    if spec.strip().lower() in ("todas", "all", ""):
        return list(range(total))
    indices = set()
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            indices.update(range(int(a) - 1, int(b)))
        elif part.isdigit():
            indices.add(int(part) - 1)
    return sorted(i for i in indices if 0 <= i < total)


# ── UI ────────────────────────────────────────────────────────────────────────

with st.expander("ℹ️ ¿Cómo obtener la API Key de Groq? (es gratis)"):
    st.markdown("""
1. Ve a [console.groq.com](https://console.groq.com) y crea una cuenta gratuita
2. En el menú izquierdo: **API Keys** → **Create API Key**
3. Copia la clave (empieza con `gsk_`) y pégala abajo

**Límites gratuitos:** 14,400 requests/día · 30 req/min — más que suficiente para 40 páginas.
    """)

api_key = st.text_input(
    "🔑 Groq API Key",
    type="password",
    placeholder="gsk_...",
    help="Gratis en console.groq.com. No se almacena.",
)

uploaded_file = st.file_uploader("📄 Sube tu PDF", type=["pdf"])

col1, col2, col3 = st.columns(3)
with col1:
    page_range = st.text_input(
        "Páginas a procesar",
        value="todas",
        placeholder="todas  ó  1-10  ó  1,3,5-8",
    )
with col2:
    page_mode = st.selectbox(
        "Tipo de páginas en el PDF",
        options=["Detectar automáticamente", "Siempre doble (landscape)", "Siempre simple"],
        help=(
            "• Doble: cada PDF contiene dos páginas del libro lado a lado\n"
            "• Simple: cada PDF es una sola página del libro\n"
            "• Automático: detecta por proporción ancho/alto"
        ),
    )
with col3:
    show_raw = st.checkbox("Mostrar texto crudo extraído", value=False)

if st.button("🚀 Procesar PDF", disabled=not (api_key and uploaded_file), type="primary"):

    api_key = api_key.strip()
    if not api_key.startswith("gsk_"):
        st.error("⛔ La API key de Groq debe comenzar con `gsk_`. Obtén la tuya en [console.groq.com](https://console.groq.com).")
        st.stop()

    client = Groq(api_key=api_key)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        tmp_pdf.write(uploaded_file.read())
        pdf_path = tmp_pdf.name

    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)
    pdf_doc.close()

    pages_to_process = parse_page_range(page_range, total_pages)
    st.info(f"Procesando {len(pages_to_process)} páginas PDF ({total_pages} total)…")

    progress = st.progress(0)
    status = st.empty()
    all_pages = []
    errors = []
    blank_count = 0
    raw_texts = {}

    total_subpages = 0
    processed = 0

    # First pass: count total sub-pages for progress bar
    for page_idx in pages_to_process:
        if page_mode == "Siempre doble (landscape)":
            total_subpages += 2
        elif page_mode == "Siempre simple":
            total_subpages += 1
        else:
            mode = detect_page_mode(pdf_path, page_idx)
            total_subpages += 2 if mode == "double" else 1

    for page_idx in pages_to_process:
        page_num = page_idx + 1

        # Decide split mode for this page
        if page_mode == "Siempre doble (landscape)":
            mode = "double"
        elif page_mode == "Siempre simple":
            mode = "single"
        else:
            mode = detect_page_mode(pdf_path, page_idx)

        halves = ["left", "right"] if mode == "double" else ["full"]

        for half in halves:
            label = (
                f"página {page_num} ({half})"
                if mode == "double"
                else f"página {page_num}"
            )
            status.text(f"Procesando {label}…")

            try:
                raw_text = extract_text_half(pdf_path, page_idx, half)
                raw_texts[label] = raw_text

                if is_text_empty(raw_text):
                    all_pages.append({
                        "_label": label,
                        "blank": True,
                        "page_header": None,
                        "poems": [],
                        "footnotes": [],
                    })
                    blank_count += 1
                else:
                    result = structure_page(client, raw_text, label)
                    result["_label"] = label
                    all_pages.append(result)

            except json.JSONDecodeError as e:
                errors.append(f"{label}: JSON inválido — {e}")
            except Exception as e:
                errors.append(f"{label}: {e}")

            processed += 1
            progress.progress(processed / total_subpages)

    os.unlink(pdf_path)
    progress.empty()
    status.empty()

    if errors:
        with st.expander(f"⚠️ {len(errors)} errores"):
            for err in errors:
                st.code(err)

    if all_pages:
        ok = len(all_pages) - len(errors)
        st.success(
            f"✅ {ok} sub-páginas procesadas · {blank_count} en blanco preservadas."
        )

        if show_raw:
            with st.expander("📄 Texto crudo por sub-página"):
                for label, text in raw_texts.items():
                    st.markdown(f"**{label}**")
                    st.code(text)

        with st.expander("🔍 Ver datos estructurados (JSON)"):
            st.json(all_pages)

        with st.spinner("Generando documento Word…"):
            docx_bytes = build_docx(all_pages)

        st.download_button(
            label="⬇️ Descargar documento Word",
            data=docx_bytes,
            file_name="poemas_extraidos.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.error("No se pudo extraer contenido de ninguna página.")
